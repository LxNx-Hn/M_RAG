"""Build template-neutral Korean and English thesis DOCX files from Markdown.

Run from the repository root with the bundled workspace Python runtime.
Institution-specific cover and approval pages are intentionally excluded because
their metadata is not present in the repository.
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[3]
PAPER_DIR = ROOT / "docs" / "PAPER"
OUTPUT_DIR = PAPER_DIR / "output"
ASSET_DIR = Path(tempfile.gettempdir()) / "mrag_thesis_build_assets"

NAVY = "17324D"
BLUE = "2563A6"
MUTED = "5A6B7B"
LIGHT = "E8EEF5"
GREEN = "EAF5EA"
GRID = "A9B7C4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font_for(lang: str, size: int, bold=False):
    path = Path(
        "C:/Windows/Fonts/malgun.ttf" if lang == "ko" else "C:/Windows/Fonts/arial.ttf"
    )
    if bold:
        bold_path = Path(
            "C:/Windows/Fonts/malgunbd.ttf"
            if lang == "ko"
            else "C:/Windows/Fonts/arialbd.ttf"
        )
        if bold_path.exists():
            path = bold_path
    return ImageFont.truetype(str(path), size)


def centered_text(draw, box, text, font, fill=NAVY):
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
    y = y1 + (y2 - y1 - (bbox[3] - bbox[1])) / 2
    draw.text((x, y), text, font=font, fill="#" + fill)


def make_diagram_assets(lang: str) -> tuple[Path, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    system = ASSET_DIR / f"system_overview_{lang}.png"
    factorial = ASSET_DIR / f"factorial_design_{lang}.png"
    title_font = font_for(lang, 34, True)
    head_font = font_for(lang, 24, True)
    body_font = font_for(lang, 19)
    small_font = font_for(lang, 16)

    labels = {
        "ko": {
            "sys_title": "M-RAG 연구 계층과 서비스 계층",
            "corpus": "논문 코퍼스\n파싱 · 절 탐지 · 청킹",
            "backbone": "고정 Paper-RAG Backbone\nBGE-M3 + BM25 + weighted RRF\nCrossEncoder reranking",
            "research": "통제된 연구 계층\nHyDE × CAD × SCD\n8설정 · 19질의 · 152답변",
            "service": "서비스 구현 계층\nFastAPI · React · SSE · 출처 표시",
            "routes": "A 단순 QA     B 절 QA     C 비교     D 인용     E 요약     F 퀴즈",
            "note": "동일 코드베이스에서 통제 실험과 A–F 논문 질의 기능을 구현한다.",
            "fac_title": "HyDE × CAD × SCD 2×2×2 조합 실험",
            "inputs": "고정 입력\n한국어 질의 19개\n영어 논문 4편",
            "factors": "세 이진 요인\nHyDE: 검색 확장\nCAD: 근거 충실도 제어\nSCD: 한국어 출력 제어",
            "matrix": "실행 행렬\n8개 설정\n152개 생성",
            "measured": "측정\nfaithfulness · answer relevancy\ncontext precision · recall\n직접 한국어 준수율",
            "deferred": "향후 평가\n질의 유형별 효과\n숫자 정확성 주석",
        },
        "en": {
            "sys_title": "M-RAG Research and Service Layers",
            "corpus": "Paper corpus\nparse · section · chunk",
            "backbone": "Fixed Paper-RAG Backbone\nBGE-M3 + BM25 + weighted RRF\nCrossEncoder reranking",
            "research": "Controlled Research Layer\nHyDE × CAD × SCD\n8 configs · 19 queries · 152 outputs",
            "service": "Service Implementation Layer\nFastAPI · React · SSE · sources",
            "routes": "A QA     B Section     C Compare     D Citation     E Summary     F Quiz",
            "note": "One codebase implements controlled experiments and A–F paper-QA functions.",
            "fac_title": "HyDE × CAD × SCD 2×2×2 Combination Experiment",
            "inputs": "Frozen inputs\n19 Korean queries\n4 English papers",
            "factors": "Three binary factors\nHyDE: retrieval reformulation\nCAD: evidence control\nSCD: Korean language control",
            "matrix": "Executed matrix\n8 configurations\n152 generations",
            "measured": "Measured\nfaithfulness · answer relevancy\ncontext precision · recall\ndirect Korean adherence",
            "deferred": "Future evaluation\nquery-type effects\nnumeric-exactness annotation",
        },
    }[lang]

    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    centered_text(d, (0, 25, 1600, 90), labels["sys_title"], title_font)
    boxes = [
        ((55, 170, 345, 340), labels["corpus"], "#F5F8FB"),
        ((430, 140, 900, 370), labels["backbone"], "#F5F8FB"),
        ((995, 120, 1545, 390), labels["research"], "#E8F2FF"),
    ]
    for box, text, fill in boxes:
        d.rounded_rectangle(box, 18, fill=fill, outline="#557A9E", width=3)
        lines = text.split("\n")
        y = box[1] + 35
        for idx, line in enumerate(lines):
            centered_text(
                d,
                (box[0], y, box[2], y + 45),
                line,
                head_font if idx == 0 else body_font,
            )
            y += 50
    d.line((345, 255, 430, 255), fill="#50677D", width=5)
    d.line((900, 255, 995, 255), fill="#50677D", width=5)
    d.rounded_rectangle(
        (200, 520, 1400, 760), 20, fill="#EEF8EE", outline="#3D7A46", width=3
    )
    centered_text(d, (200, 545, 1400, 605), labels["service"].split("\n")[0], head_font)
    centered_text(d, (200, 605, 1400, 655), labels["service"].split("\n")[1], body_font)
    centered_text(d, (220, 665, 1380, 725), labels["routes"], body_font)
    centered_text(d, (0, 810, 1600, 860), labels["note"], small_font, MUTED)
    img.save(system, dpi=(180, 180))

    img = Image.new("RGB", (1600, 850), "white")
    d = ImageDraw.Draw(img)
    centered_text(d, (0, 25, 1600, 90), labels["fac_title"], title_font)
    boxes = [
        ((60, 140, 390, 350), labels["inputs"], "#F5F8FB"),
        ((490, 115, 1110, 380), labels["factors"], "#E8F2FF"),
        ((1210, 140, 1540, 350), labels["matrix"], "#EAF5EA"),
    ]
    for box, text, fill in boxes:
        d.rounded_rectangle(box, 18, fill=fill, outline="#557A9E", width=3)
        lines = text.split("\n")
        y = box[1] + 35
        for idx, line in enumerate(lines):
            centered_text(
                d,
                (box[0], y, box[2], y + 45),
                line,
                head_font if idx == 0 else body_font,
            )
            y += 50
    d.line((390, 245, 490, 245), fill="#50677D", width=5)
    d.line((1110, 245, 1210, 245), fill="#50677D", width=5)
    for box, key, fill in [
        ((150, 500, 760, 740), "measured", "#F5F8FB"),
        ((840, 500, 1450, 740), "deferred", "#F5F8FB"),
    ]:
        d.rounded_rectangle(box, 18, fill=fill, outline="#557A9E", width=3)
        y = box[1] + 30
        for idx, line in enumerate(labels[key].split("\n")):
            centered_text(
                d,
                (box[0], y, box[2], y + 42),
                line,
                head_font if idx == 0 else body_font,
            )
            y += 47
    img.save(factorial, dpi=(180, 180))
    return system, factorial


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def configure_document(doc: Document, lang: str) -> None:
    # narrative_proposal preset with named academic_a4 override.
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    base_font = "Malgun Gothic" if lang == "ko" else "Arial"
    for style_name, size, color, before, after in [
        ("Normal", 10.5, "000000", 0, 7),
        ("Heading 1", 16, NAVY, 16, 9),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = base_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), base_font)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.35 if style_name == "Normal" else 1.15
        if style_name != "Normal":
            style.font.bold = True
            style.paragraph_format.keep_with_next = True
    normal = doc.styles["Normal"]
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = base_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), base_font)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.75)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    header = section.header.paragraphs[0]
    header.text = "M-RAG Thesis" if lang == "en" else "M-RAG 학위논문 원고"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = base_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), base_font)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_cover(doc: Document, title: str, lang: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Malgun Gothic" if lang == "ko" else "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table(doc: Document, rows: list[list[str]], lang: str) -> None:
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(cols):
            cells[cidx].text = clean_inline(row[cidx] if cidx < len(row) else "")
            if ridx == 0:
                set_cell_shading(cells[cidx], LIGHT)
            for paragraph in cells[cidx].paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if cidx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Malgun Gothic" if lang == "ko" else "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
                    run.font.size = Pt(8.5 if cols >= 5 else 9)
                    if ridx == 0:
                        run.bold = True
    set_repeat_table_header(table.rows[0])
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2200, 3580, 3580]
    elif cols == 4:
        widths = [2200, 2380, 2380, 2400]
    elif cols == 5:
        widths = [2600, 1690, 1690, 1690, 1690]
    else:
        widths = [2350, 1402, 1402, 1402, 1402, 1402]
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build(markdown_path: Path, output_path: Path, lang: str) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    title = clean_inline(lines[0].lstrip("# "))
    system_img, factorial_img = make_diagram_assets(lang)
    doc = Document()
    configure_document(doc, lang)
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = title
    doc.core_properties.subject = "M-RAG thesis manuscript"
    add_cover(doc, title, lang)
    code_mode = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    para_lines: list[str] = []

    def flush_para():
        nonlocal para_lines
        if para_lines:
            text = clean_inline(" ".join(part.strip() for part in para_lines))
            if text:
                doc.add_paragraph(text)
            para_lines = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            rows = [
                row
                for row in table_rows
                if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)
            ]
            if rows:
                add_table(doc, rows, lang)
            table_rows = []

    for line in lines[1:]:
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_para()
            flush_table()
            if code_mode:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.05
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F4F7")
                p_pr.append(shd)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(8.5)
                code_lines = []
            code_mode = not code_mode
            continue
        if code_mode:
            code_lines.append(line)
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_para()
            table_rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
            continue
        flush_table()
        if not stripped:
            flush_para()
            continue
        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_para()
            src = image_match.group(2)
            image = system_img if "system_overview" in src else factorial_img
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_run = p.add_run()
            picture_run.add_picture(str(image), width=Inches(6.2))
            alt_text = clean_inline(image_match.group(1))
            for doc_pr in picture_run._r.xpath(".//wp:docPr"):
                doc_pr.set("descr", alt_text)
                doc_pr.set("title", alt_text)
            continue
        if stripped.startswith("### "):
            flush_para()
            doc.add_heading(clean_inline(stripped[4:]), level=2)
            continue
        if stripped.startswith("## "):
            flush_para()
            doc.add_heading(clean_inline(stripped[3:]), level=1)
            continue
        if stripped.startswith("# "):
            flush_para()
            doc.add_heading(clean_inline(stripped[2:]), level=1)
            continue
        if stripped.startswith("- "):
            flush_para()
            doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
            continue
        if re.match(r"^\d+\. ", stripped):
            flush_para()
            doc.add_paragraph(
                clean_inline(re.sub(r"^\d+\. ", "", stripped)), style="List Number"
            )
            continue
        if stripped.startswith("> "):
            flush_para()
            p = doc.add_paragraph(clean_inline(stripped[2:]))
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.7)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), GREEN)
            p_pr.append(shd)
            for run in p.runs:
                run.bold = True
            continue
        para_lines.append(stripped)
    flush_para()
    flush_table()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    build(PAPER_DIR / "THESIS_KO.md", OUTPUT_DIR / "M_RAG_THESIS_KO.docx", "ko")
    build(PAPER_DIR / "THESIS.md", OUTPUT_DIR / "M_RAG_THESIS_EN.docx", "en")
    print("Built Korean and English submission manuscripts.")
